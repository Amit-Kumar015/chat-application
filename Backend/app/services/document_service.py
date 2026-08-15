import os
import tempfile
import logging
from typing import List
from fastapi import UploadFile, HTTPException
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, WebBaseLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxReader

logger = logging.getLogger(__name__)

text_splitter = RecursiveCharacterTextSplitter(
  chunk_size=1000,
  chunk_overlap=200,
  length_function=len
)

async def parse_uploaded_file(file: UploadFile) -> List[Document]:
  filename = file.filename or "uploaded_file"
  extension = os.path.splitext(filename)[1].lower()
  
  temp_file_path = None
  try:
    with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as tmp:
      content = await file.read()
      tmp.write(content)
      temp_file_path = tmp.name
    
    raw_documents: Document = []
    
    if extension == ".pdf":
      loader = PyPDFLoader(temp_file_path)
      raw_documents = loader.load()
    elif extension in [".doc", ".docx"]:
      doc = DocxReader(temp_file_path)
      full_text = "\n".join([p.text for p in doc.paragraphs if p.text.split()])
      raw_documents = [Document(page_content=full_text, metadata={"source": filename})]
    elif extension == ".txt":
      with open(temp_file_path, "r", encoding="utf-8") as f:
        full_text = f.read()
      raw_documents = [Document(page_content=full_text, metadata={"source": filename})]
    else:
      raise HTTPException(status_code=400, detail=f"Unsupported file format '{extension}'. Allowed: .pdf, .docx, .txt")
    
  except HTTPException:
    raise
  except Exception as e:
    logger.error(f"Error processing file {filename}: {e}")
    raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")
  finally:
    if temp_file_path and os.path.exists(temp_file_path):
      os.remove(temp_file_path)
      logger.info(f"Temporary file deleted: {temp_file_path}")
      
  chunks = text_splitter.split_documents(raw_documents)
  
  for chunk in chunks:
    chunk.metadata["source"] = filename
  
  return chunks

async def parse_url(url: str) -> List[Document]:
  try:
    loader = WebBaseLoader(url)
    raw_documents = loader.load()
    chunks = text_splitter.split_documents(raw_documents)
    
    for chunk in chunks:
      chunk.metadata["source"] = url
      
    return chunks
  
  except Exception as e:
    logger.error(f"Error scraping URL {url}: {e}")
    raise HTTPException(status_code=500, detail=f"Failed to scrape URL: {str(e)}")